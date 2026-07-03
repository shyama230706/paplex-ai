import streamlit as st
import os
import json
import re
import hashlib
from io import BytesIO

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PapLex AI",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ── Custom CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }

    .stApp {
        background: linear-gradient(135deg, #0f0c29, #302b63, #24243e);
        color: #e2e8f0;
    }

    section[data-testid="stSidebar"] {
        background: rgba(15, 12, 41, 0.95) !important;
        border-right: 1px solid rgba(139, 92, 246, 0.3);
    }

    .main-header {
        background: linear-gradient(135deg, #4c1d95, #6d28d9, #7c3aed);
        border-radius: 20px;
        padding: 2.5rem;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 20px 60px rgba(124, 58, 237, 0.4);
    }

    .main-header h1 {
        font-size: 3rem;
        font-weight: 700;
        color: white;
        margin: 0;
    }

    .main-header p {
        color: rgba(255,255,255,0.85);
        font-size: 1.1rem;
        margin: 0.5rem 0 0 0;
    }

    .feature-tags {
        display: flex;
        justify-content: center;
        gap: 0.5rem;
        flex-wrap: wrap;
        margin-top: 1rem;
    }

    .feature-tag {
        background: rgba(255,255,255,0.15);
        color: white;
        padding: 0.3rem 0.8rem;
        border-radius: 20px;
        font-size: 0.8rem;
        border: 1px solid rgba(255,255,255,0.2);
    }

    .chat-message-user {
        background: linear-gradient(135deg, #4c1d95, #6d28d9);
        border-radius: 15px 15px 5px 15px;
        padding: 1rem 1.2rem;
        margin: 0.5rem 0;
        color: white;
        margin-left: 20%;
        box-shadow: 0 4px 15px rgba(109, 40, 217, 0.3);
    }

    .chat-message-bot {
        background: rgba(30, 27, 75, 0.8);
        border: 1px solid rgba(139, 92, 246, 0.3);
        border-radius: 15px 15px 15px 5px;
        padding: 1rem 1.2rem;
        margin: 0.5rem 0;
        color: #e2e8f0;
        margin-right: 10%;
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    }

    .stButton > button {
        background: linear-gradient(135deg, #6d28d9, #7c3aed) !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        padding: 0.6rem 1.5rem !important;
        font-weight: 600 !important;
        transition: all 0.3s !important;
        box-shadow: 0 4px 15px rgba(109, 40, 217, 0.3) !important;
    }

    .stButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 25px rgba(109, 40, 217, 0.5) !important;
    }

    .info-card {
        background: rgba(30, 27, 75, 0.6);
        border: 1px solid rgba(139, 92, 246, 0.3);
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
    }

    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        background: rgba(30, 27, 75, 0.8) !important;
        border: 1px solid rgba(139, 92, 246, 0.4) !important;
        color: #e2e8f0 !important;
        border-radius: 10px !important;
    }

    .stSelectbox > div > div {
        background: rgba(30, 27, 75, 0.8) !important;
        border: 1px solid rgba(139, 92, 246, 0.4) !important;
        color: #e2e8f0 !important;
        border-radius: 10px !important;
    }

    div[data-testid="stFileUploader"] {
        background: rgba(30, 27, 75, 0.5) !important;
        border: 2px dashed rgba(139, 92, 246, 0.5) !important;
        border-radius: 15px !important;
        padding: 1rem !important;
    }

    .metric-card {
        background: rgba(30, 27, 75, 0.6);
        border: 1px solid rgba(139, 92, 246, 0.3);
        border-radius: 12px;
        padding: 1rem;
        text-align: center;
    }

    hr {
        border-color: rgba(139, 92, 246, 0.3) !important;
    }
</style>
""", unsafe_allow_html=True)

# ── Imports (lazy loaded) ──────────────────────────────────────────────────────
try:
    from groq import Groq
    from langchain_groq import ChatGroq
    from langchain_community.vectorstores import FAISS
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain.schema import Document
    import pandas as pd
    import plotly.express as px
    import plotly.graph_objects as go
    IMPORTS_OK = True
except ImportError as e:
    IMPORTS_OK = False
    IMPORT_ERROR = str(e)

# ── API Key ────────────────────────────────────────────────────────────────────
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")

# ── Session State ──────────────────────────────────────────────────────────────
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "doc_names" not in st.session_state:
    st.session_state.doc_names = []
if "page" not in st.session_state:
    st.session_state.page = "Chat"

# ── Helper Functions ───────────────────────────────────────────────────────────
def extract_text_from_file(uploaded_file):
    """Extract text from various file formats."""
    name = uploaded_file.name.lower()
    text = ""

    try:
        if name.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(uploaded_file.read()))
            for page in reader.pages:
                text += page.extract_text() or ""

        elif name.endswith(".txt"):
            text = uploaded_file.read().decode("utf-8", errors="ignore")

        elif name.endswith(".csv"):
            import pandas as pd
            df = pd.read_csv(uploaded_file)
            text = df.to_string()

        elif name.endswith((".xlsx", ".xls")):
            import pandas as pd
            df = pd.read_excel(uploaded_file)
            text = df.to_string()

        elif name.endswith(".json"):
            data = json.load(uploaded_file)
            text = json.dumps(data, indent=2)

        elif name.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(BytesIO(uploaded_file.read()))
            for para in doc.paragraphs:
                text += para.text + "\n"

        else:
            text = uploaded_file.read().decode("utf-8", errors="ignore")

    except Exception as e:
        text = f"Error reading file: {str(e)}"

    return text


def build_vectorstore(texts, api_key):
    """Build FAISS vectorstore using lightweight embeddings."""
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
    docs = []
    for t in texts:
        chunks = splitter.split_text(t)
        for chunk in chunks:
            docs.append(Document(page_content=chunk))

    # Use lightweight TF-IDF style via simple hash embeddings for free tier
    # Actually use a very small model
    try:
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True}
        )
        vectorstore = FAISS.from_documents(docs, embeddings)
    except Exception:
        # Fallback: simple keyword search without embeddings
        vectorstore = None
        st.session_state.raw_docs = docs

    return vectorstore


def get_groq_response(prompt, context, language, api_key, history=None):
    """Get response from Groq LLM."""
    client = Groq(api_key=api_key)

    lang_instruction = f"Always respond in {language}." if language != "English" else ""

    system_prompt = f"""You are PapLex AI, an intelligent document assistant.
Answer questions based on the provided document context.
Be helpful, accurate, and concise.
{lang_instruction}
If the answer is not in the context, say so honestly."""

    messages = [{"role": "system", "content": system_prompt}]

    if history:
        for h in history[-4:]:
            messages.append({"role": "user", "content": h["user"]})
            messages.append({"role": "assistant", "content": h["bot"]})

    messages.append({
        "role": "user",
        "content": f"Context:\n{context}\n\nQuestion: {prompt}"
    })

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=messages,
        max_tokens=1024,
        temperature=0.7
    )

    return response.choices[0].message.content


def search_context(query, vectorstore, raw_docs=None):
    """Search for relevant context."""
    if vectorstore:
        docs = vectorstore.similarity_search(query, k=4)
        return "\n\n".join([d.page_content for d in docs])
    elif raw_docs:
        # Simple keyword fallback
        query_words = query.lower().split()
        scored = []
        for doc in raw_docs:
            score = sum(1 for w in query_words if w in doc.page_content.lower())
            scored.append((score, doc.page_content))
        scored.sort(reverse=True)
        return "\n\n".join([c for _, c in scored[:4]])
    return ""


# ── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 🧠 PapLex AI")
    st.markdown('<div style="background: linear-gradient(135deg, #6d28d9, #7c3aed); color: white; padding: 0.3rem 0.8rem; border-radius: 20px; font-size: 0.75rem; display: inline-block; margin-bottom: 1rem;">Powered by LLaMA 3.3</div>', unsafe_allow_html=True)
    st.divider()

    # Navigation
    st.markdown("### 🧭 Navigation")
    pages = ["💬 Chat", "🔄 Convert Files", "📊 Visualize"]
    for p in pages:
        if st.button(p, use_container_width=True, key=f"nav_{p}"):
            st.session_state.page = p.split(" ", 1)[1]

    st.divider()

    # Language
    st.markdown("### 🌍 Language")
    language = st.selectbox(
        "Answer language:",
        ["English", "Hindi", "Spanish", "French", "German", "Japanese", "Chinese", "Arabic"],
        label_visibility="collapsed"
    )

    st.divider()

    # File Upload
    st.markdown("### 📁 Upload Files")
    st.caption("PDF, CSV, Excel, JSON, TXT, Word")

    uploaded_files = st.file_uploader(
        "Upload documents",
        type=["pdf", "txt", "csv", "xlsx", "xls", "json", "docx"],
        accept_multiple_files=True,
        label_visibility="collapsed"
    )

    if uploaded_files:
        if st.button("🚀 Process Documents", use_container_width=True):
            with st.spinner("Processing documents..."):
                texts = []
                names = []
                for f in uploaded_files:
                    text = extract_text_from_file(f)
                    if text:
                        texts.append(text)
                        names.append(f.name)

                if texts:
                    vs = build_vectorstore(texts, GROQ_API_KEY)
                    st.session_state.vectorstore = vs
                    st.session_state.doc_names = names
                    st.session_state.raw_texts = texts
                    st.success(f"✅ {len(names)} file(s) processed!")
                else:
                    st.error("Could not extract text from files.")

    if st.session_state.doc_names:
        st.divider()
        st.markdown("### 📄 Loaded Documents")
        for name in st.session_state.doc_names:
            st.markdown(f"✅ `{name}`")

        if st.button("🗑️ Clear All", use_container_width=True):
            st.session_state.vectorstore = None
            st.session_state.doc_names = []
            st.session_state.chat_history = []
            st.session_state.raw_texts = []
            st.rerun()

    st.divider()

    # API Key input (if not set via env)
    if not GROQ_API_KEY:
        st.markdown("### 🔑 API Key")
        user_key = st.text_input("Groq API Key", type="password", placeholder="gsk_...")
        if user_key:
            GROQ_API_KEY = user_key


# ── Main Content ───────────────────────────────────────────────────────────────
current_page = st.session_state.get("page", "Chat")

# ── CHAT PAGE ──────────────────────────────────────────────────────────────────
if "Chat" in current_page or current_page == "Chat":

    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🧠 PapLex AI</h1>
        <p>Your Intelligent Document Assistant</p>
        <div class="feature-tags">
            <span class="feature-tag">💬 Chat</span>
            <span class="feature-tag">🔄 Convert</span>
            <span class="feature-tag">📊 Visualize</span>
            <span class="feature-tag">🌍 Multilingual</span>
            <span class="feature-tag">📁 6 Formats</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.doc_names:
        # Welcome screen
        st.markdown("""
        <div class="info-card" style="text-align: center;">
            <h2>👋 Welcome to PapLex AI</h2>
            <h4>Your Intelligent Document Assistant</h4>
            <p style="color: rgba(255,255,255,0.7);">👆 Upload files from sidebar OR use Convert Files tab!</p>
        </div>
        """, unsafe_allow_html=True)

        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("""
            <div class="info-card" style="text-align:center;">
                <h3>📄 Supported Formats</h3>
                <p>PDF • CSV • Excel<br>JSON • TXT • Word</p>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown("""
            <div class="info-card" style="text-align:center;">
                <h3>🤖 AI Features</h3>
                <p>Q&A • Summary • Quiz<br>NER • Comparison</p>
            </div>
            """, unsafe_allow_html=True)
        with col3:
            st.markdown("""
            <div class="info-card" style="text-align:center;">
                <h3>🌍 Languages</h3>
                <p>English • Hindi • Spanish<br>French • German • More</p>
            </div>
            """, unsafe_allow_html=True)

        # Smart Actions (no doc needed)
        st.divider()
        st.markdown("### ⚡ Quick Actions (No Document Needed)")
        col1, col2, col3, col4, col5 = st.columns(5)

        action = None
        with col1:
            if st.button("📝 Summarize", use_container_width=True):
                action = "summarize"
        with col2:
            if st.button("❓ Quiz Me", use_container_width=True):
                action = "quiz"
        with col3:
            if st.button("🏷️ Extract Entities", use_container_width=True):
                action = "ner"
        with col4:
            if st.button("💡 Suggest Questions", use_container_width=True):
                action = "suggest"
        with col5:
            if st.button("🔍 Compare Docs", use_container_width=True):
                action = "compare"

    else:
        # Document loaded — show chat
        st.markdown(f"### 💬 Chat with your documents")
        st.caption(f"📄 Loaded: {', '.join(st.session_state.doc_names)}")

        # Smart Actions
        st.markdown("#### ⚡ Smart Actions")
        col1, col2, col3, col4, col5 = st.columns(5)
        action = None
        with col1:
            if st.button("📝 Summarize", use_container_width=True):
                action = "summarize"
        with col2:
            if st.button("❓ Quiz Me", use_container_width=True):
                action = "quiz"
        with col3:
            if st.button("🏷️ Entities", use_container_width=True):
                action = "ner"
        with col4:
            if st.button("💡 Suggest", use_container_width=True):
                action = "suggest"
        with col5:
            if st.button("🔍 Compare", use_container_width=True):
                action = "compare"

        # Handle smart actions
        if action and GROQ_API_KEY:
            raw_text = " ".join(getattr(st.session_state, 'raw_texts', []))[:3000]
            prompts = {
                "summarize": "Please provide a comprehensive summary of the document(s).",
                "quiz": "Generate 5 multiple choice quiz questions from the document content with answers.",
                "ner": "Extract and list all named entities (persons, organizations, locations, dates, numbers) from the document.",
                "suggest": "Suggest 8 insightful questions that someone could ask about this document.",
                "compare": "If there are multiple documents, compare and contrast them. Otherwise, analyze the key themes."
            }
            with st.spinner("🤖 AI is thinking..."):
                response = get_groq_response(prompts[action], raw_text, language, GROQ_API_KEY)
                st.session_state.chat_history.append({
                    "user": f"[{action.upper()}]",
                    "bot": response
                })
            st.rerun()

        st.divider()

        # Chat history
        if st.session_state.chat_history:
            for chat in st.session_state.chat_history:
                st.markdown(f'<div class="chat-message-user">👤 {chat["user"]}</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="chat-message-bot">🧠 {chat["bot"]}</div>', unsafe_allow_html=True)

        # Chat input
        st.divider()
        col1, col2 = st.columns([5, 1])
        with col1:
            user_input = st.text_input(
                "Ask anything about your documents...",
                placeholder="e.g. What are the main findings?",
                label_visibility="collapsed",
                key="chat_input"
            )
        with col2:
            send = st.button("Send 🚀", use_container_width=True)

        if (send or user_input) and user_input and GROQ_API_KEY:
            with st.spinner("🤖 Thinking..."):
                context = search_context(
                    user_input,
                    st.session_state.vectorstore,
                    getattr(st.session_state, 'raw_docs', None)
                )
                if not context:
                    context = " ".join(getattr(st.session_state, 'raw_texts', []))[:3000]

                response = get_groq_response(
                    user_input, context, language, GROQ_API_KEY,
                    st.session_state.chat_history
                )
                st.session_state.chat_history.append({
                    "user": user_input,
                    "bot": response
                })
            st.rerun()

        if not GROQ_API_KEY:
            st.warning("⚠️ Please enter your Groq API key in the sidebar.")

        if st.session_state.chat_history:
            if st.button("🗑️ Clear Chat"):
                st.session_state.chat_history = []
                st.rerun()


# ── CONVERT FILES PAGE ─────────────────────────────────────────────────────────
elif "Convert" in current_page:
    st.markdown("""
    <div class="main-header">
        <h1>🔄 File Converter</h1>
        <p>Convert between PDF, CSV, Excel, JSON, TXT, Word</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        convert_file = st.file_uploader(
            "Upload file to convert",
            type=["pdf", "txt", "csv", "xlsx", "json", "docx"],
            key="convert_upload"
        )
    with col2:
        target_format = st.selectbox(
            "Convert to:",
            ["TXT", "CSV", "JSON", "Excel (XLSX)"]
        )

    if convert_file and st.button("🔄 Convert Now", use_container_width=True):
        with st.spinner("Converting..."):
            text = extract_text_from_file(convert_file)

            if target_format == "TXT":
                result = text.encode("utf-8")
                st.download_button("⬇️ Download TXT", result,
                                   file_name=f"{convert_file.name}.txt",
                                   mime="text/plain")

            elif target_format == "JSON":
                data = {"filename": convert_file.name, "content": text}
                result = json.dumps(data, indent=2).encode("utf-8")
                st.download_button("⬇️ Download JSON", result,
                                   file_name=f"{convert_file.name}.json",
                                   mime="application/json")

            elif target_format == "CSV":
                lines = [l for l in text.split("\n") if l.strip()]
                import pandas as pd
                df = pd.DataFrame({"content": lines})
                result = df.to_csv(index=False).encode("utf-8")
                st.download_button("⬇️ Download CSV", result,
                                   file_name=f"{convert_file.name}.csv",
                                   mime="text/csv")

            elif target_format == "Excel (XLSX)":
                lines = [l for l in text.split("\n") if l.strip()]
                import pandas as pd
                df = pd.DataFrame({"content": lines})
                buf = BytesIO()
                df.to_excel(buf, index=False)
                st.download_button("⬇️ Download Excel", buf.getvalue(),
                                   file_name=f"{convert_file.name}.xlsx",
                                   mime="application/vnd.ms-excel")

            st.success("✅ Conversion complete!")


# ── VISUALIZE PAGE ─────────────────────────────────────────────────────────────
elif "Visualize" in current_page:
    st.markdown("""
    <div class="main-header">
        <h1>📊 Data Visualizer</h1>
        <p>Upload CSV or Excel to create interactive charts</p>
    </div>
    """, unsafe_allow_html=True)

    viz_file = st.file_uploader(
        "Upload CSV or Excel file",
        type=["csv", "xlsx", "xls"],
        key="viz_upload"
    )

    if viz_file:
        try:
            import pandas as pd
            import plotly.express as px

            if viz_file.name.endswith(".csv"):
                df = pd.read_csv(viz_file)
            else:
                df = pd.read_excel(viz_file)

            st.success(f"✅ Loaded: {df.shape[0]} rows × {df.shape[1]} columns")
            st.dataframe(df.head(10), use_container_width=True)

            st.divider()
            st.markdown("### 📈 Create Chart")

            col1, col2, col3 = st.columns(3)
            with col1:
                chart_type = st.selectbox("Chart Type", ["Bar", "Line", "Scatter", "Pie", "Histogram", "Heatmap"])
            with col2:
                x_col = st.selectbox("X Axis", df.columns.tolist())
            with col3:
                y_col = st.selectbox("Y Axis", df.columns.tolist())

            if st.button("📊 Generate Chart", use_container_width=True):
                fig = None
                if chart_type == "Bar":
                    fig = px.bar(df, x=x_col, y=y_col, title=f"{y_col} by {x_col}",
                                 color_discrete_sequence=["#7c3aed"])
                elif chart_type == "Line":
                    fig = px.line(df, x=x_col, y=y_col, title=f"{y_col} over {x_col}",
                                  color_discrete_sequence=["#7c3aed"])
                elif chart_type == "Scatter":
                    fig = px.scatter(df, x=x_col, y=y_col, title=f"{x_col} vs {y_col}",
                                     color_discrete_sequence=["#7c3aed"])
                elif chart_type == "Pie":
                    fig = px.pie(df, names=x_col, values=y_col, title=f"{y_col} Distribution",
                                 color_discrete_sequence=px.colors.sequential.Purples)
                elif chart_type == "Histogram":
                    fig = px.histogram(df, x=x_col, title=f"Distribution of {x_col}",
                                       color_discrete_sequence=["#7c3aed"])
                elif chart_type == "Heatmap":
                    numeric_df = df.select_dtypes(include='number')
                    fig = px.imshow(numeric_df.corr(), title="Correlation Heatmap",
                                    color_continuous_scale="Purples")

                if fig:
                    fig.update_layout(
                        paper_bgcolor="rgba(0,0,0,0)",
                        plot_bgcolor="rgba(30,27,75,0.5)",
                        font_color="#e2e8f0"
                    )
                    st.plotly_chart(fig, use_container_width=True)

        except Exception as e:
            st.error(f"Error loading file: {str(e)}")

# ── Footer ─────────────────────────────────────────────────────────────────────
st.divider()
st.markdown("""
<div style="text-align: center; color: rgba(255,255,255,0.4); font-size: 0.8rem; padding: 1rem;">
    🧠 PapLex AI • Built by <strong style="color: #7c3aed;">Shyama Mishra</strong> •
    Powered by LLaMA 3.3 70B • Groq API • FAISS • LangChain
</div>
""", unsafe_allow_html=True)
